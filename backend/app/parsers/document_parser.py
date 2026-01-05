"""
Document parser for user-uploaded files
Supports: PDF, DOCX, TXT, HTML
"""
import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from pathlib import Path
import chardet
import magic

# PDF parsing
from PyPDF2 import PdfReader

# DOCX parsing
from docx import Document as DocxDocument

# HTML parsing
from bs4 import BeautifulSoup


@dataclass
class Section:
    """Represents a section in the document"""
    level: int  # 1 = top level, 2 = subsection, etc.
    title: str
    start_char: int  # Character position where section starts
    end_char: int    # Character position where section ends
    content: str     # Section text content
    page: Optional[int] = None  # Page number (for PDFs)


@dataclass
class ParsedDocument:
    """Result of parsing a document"""
    text: str  # Full normalized text
    outline: List[Section]  # Document sections
    snippets: List[str]  # Text snippets for preview
    page_map: Optional[Dict[int, Dict[str, int]]] = None  # Page -> char positions (PDFs only)
    warnings: List[str] = field(default_factory=list)  # Parsing warnings
    metadata: Dict[str, Any] = field(default_factory=dict)  # Document metadata
    confidence_score: float = 1.0  # 0.0-1.0 extraction confidence


class DocumentParser:
    """Main document parser class"""

    SUPPORTED_FORMATS = ['pdf', 'docx', 'txt', 'html', 'htm']
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    def __init__(self):
        self.magic = magic.Magic(mime=True)

    def detect_format(self, file_bytes: bytes, filename: str) -> str:
        """
        Detect document format from file bytes and filename

        Args:
            file_bytes: File content as bytes
            filename: Original filename

        Returns:
            Format string: 'pdf', 'docx', 'txt', 'html'

        Raises:
            ValueError: If format is unsupported
        """
        # Try MIME type detection first
        mime_type = self.magic.from_buffer(file_bytes)

        if 'pdf' in mime_type.lower():
            return 'pdf'
        elif 'wordprocessingml' in mime_type.lower() or 'msword' in mime_type.lower():
            return 'docx'
        elif 'html' in mime_type.lower():
            return 'html'
        elif 'text' in mime_type.lower():
            # Check extension to differentiate HTML from plain text
            ext = Path(filename).suffix.lower()
            if ext in ['.html', '.htm']:
                return 'html'
            return 'txt'

        # Fallback to file extension
        ext = Path(filename).suffix.lower().lstrip('.')
        if ext in self.SUPPORTED_FORMATS:
            return ext

        raise ValueError(f"Unsupported file format: {mime_type} ({filename})")

    def parse(self, file_bytes: bytes, filename: str, format_hint: Optional[str] = None) -> ParsedDocument:
        """
        Parse a document and extract structured content

        Args:
            file_bytes: File content as bytes
            filename: Original filename
            format_hint: Optional format override ('pdf', 'docx', 'txt', 'html')

        Returns:
            ParsedDocument with extracted content

        Raises:
            ValueError: If file is too large or unsupported
        """
        # Check file size
        if len(file_bytes) > self.MAX_FILE_SIZE:
            raise ValueError(f"File too large: {len(file_bytes)} bytes (max {self.MAX_FILE_SIZE})")

        # Detect format
        doc_format = format_hint or self.detect_format(file_bytes, filename)

        # Parse based on format
        if doc_format == 'pdf':
            return self.parse_pdf(file_bytes)
        elif doc_format == 'docx':
            return self.parse_docx(file_bytes)
        elif doc_format == 'txt':
            return self.parse_txt(file_bytes)
        elif doc_format in ['html', 'htm']:
            return self.parse_html(file_bytes)
        else:
            raise ValueError(f"Unsupported format: {doc_format}")

    def parse_pdf(self, file_bytes: bytes) -> ParsedDocument:
        """Parse PDF file"""
        from io import BytesIO

        warnings = []
        metadata = {}

        try:
            pdf = PdfReader(BytesIO(file_bytes))
            metadata['pages'] = len(pdf.pages)
            metadata['format'] = 'pdf'

            # Extract text from all pages
            pages_text = []
            page_map = {}
            char_offset = 0

            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if not page_text.strip():
                        warnings.append(f"Page {page_num}: No text extracted (may be scanned image)")
                        page_text = f"[Page {page_num}: No text content]\n"

                    pages_text.append(page_text)

                    # Track page boundaries in text
                    page_start = char_offset
                    page_end = char_offset + len(page_text)
                    page_map[page_num] = {'start': page_start, 'end': page_end}
                    char_offset = page_end

                except Exception as e:
                    warnings.append(f"Page {page_num}: Error extracting text - {str(e)}")
                    pages_text.append(f"[Page {page_num}: Error extracting text]\n")

            full_text = "\n".join(pages_text)

            # Extract outline (sections)
            outline = self.extract_outline(full_text, page_map)

            # Generate snippets
            snippets = self.generate_snippets(full_text, num_snippets=5)

            # Calculate confidence
            confidence = self._calculate_confidence(full_text, warnings, len(pdf.pages))

            return ParsedDocument(
                text=full_text,
                outline=outline,
                snippets=snippets,
                page_map=page_map,
                warnings=warnings,
                metadata=metadata,
                confidence_score=confidence
            )

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    def parse_docx(self, file_bytes: bytes) -> ParsedDocument:
        """Parse DOCX file"""
        from io import BytesIO

        warnings = []
        metadata = {'format': 'docx'}

        try:
            doc = DocxDocument(BytesIO(file_bytes))

            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            full_text = "\n".join(paragraphs)
            metadata['paragraphs'] = len(paragraphs)

            if not full_text.strip():
                warnings.append("No text content found in document")
                full_text = "[Empty document]"

            # Extract outline
            outline = self.extract_outline(full_text)

            # Generate snippets
            snippets = self.generate_snippets(full_text, num_snippets=5)

            # Calculate confidence
            confidence = self._calculate_confidence(full_text, warnings)

            return ParsedDocument(
                text=full_text,
                outline=outline,
                snippets=snippets,
                warnings=warnings,
                metadata=metadata,
                confidence_score=confidence
            )

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def parse_txt(self, file_bytes: bytes) -> ParsedDocument:
        """Parse plain text file"""
        warnings = []
        metadata = {'format': 'txt'}

        try:
            # Detect encoding
            detection = chardet.detect(file_bytes)
            encoding = detection['encoding'] or 'utf-8'
            confidence_enc = detection['confidence']

            if confidence_enc < 0.7:
                warnings.append(f"Low encoding detection confidence: {confidence_enc:.2f}")

            # Decode text
            full_text = file_bytes.decode(encoding, errors='replace')
            metadata['encoding'] = encoding
            metadata['lines'] = len(full_text.splitlines())

            if not full_text.strip():
                warnings.append("Empty text file")
                full_text = "[Empty document]"

            # Extract outline
            outline = self.extract_outline(full_text)

            # Generate snippets
            snippets = self.generate_snippets(full_text, num_snippets=5)

            # Calculate confidence
            confidence = self._calculate_confidence(full_text, warnings)

            return ParsedDocument(
                text=full_text,
                outline=outline,
                snippets=snippets,
                warnings=warnings,
                metadata=metadata,
                confidence_score=confidence
            )

        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {str(e)}")

    def parse_html(self, file_bytes: bytes) -> ParsedDocument:
        """Parse HTML file"""
        warnings = []
        metadata = {'format': 'html'}

        try:
            # Detect encoding
            detection = chardet.detect(file_bytes)
            encoding = detection['encoding'] or 'utf-8'

            # Decode HTML
            html_content = file_bytes.decode(encoding, errors='replace')

            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, 'lxml')

            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()

            # Extract text
            full_text = soup.get_text(separator='\n', strip=True)
            metadata['encoding'] = encoding

            if not full_text.strip():
                warnings.append("No text content found in HTML")
                full_text = "[Empty document]"

            # Extract outline (from headings)
            outline = self.extract_outline_from_html(soup)

            # Generate snippets
            snippets = self.generate_snippets(full_text, num_snippets=5)

            # Calculate confidence
            confidence = self._calculate_confidence(full_text, warnings)

            return ParsedDocument(
                text=full_text,
                outline=outline,
                snippets=snippets,
                warnings=warnings,
                metadata=metadata,
                confidence_score=confidence
            )

        except Exception as e:
            raise ValueError(f"Failed to parse HTML: {str(e)}")

    def extract_outline(self, text: str, page_map: Optional[Dict] = None) -> List[Section]:
        """
        Extract document outline from text

        Detects section headings using heuristics:
        - All caps lines
        - Numbered sections (1., 1.1, etc.)
        - Common section keywords
        """
        sections = []
        lines = text.split('\n')
        char_offset = 0

        # Patterns for section detection
        numbered_pattern = re.compile(r'^(\d+\.)+\s+(.+)')
        roman_pattern = re.compile(r'^[IVXLCDM]+\.\s+(.+)')
        section_keywords = ['SECTION', 'ARTICLE', 'CHAPTER', 'PART', 'TITLE', 'CLAUSE']

        for i, line in enumerate(lines):
            line_stripped = line.strip()

            if not line_stripped:
                char_offset += len(line) + 1
                continue

            is_section = False
            level = 1
            title = line_stripped

            # Check for numbered sections (1., 1.1., etc.)
            match = numbered_pattern.match(line_stripped)
            if match:
                is_section = True
                level = line_stripped.count('.')
                title = match.group(2).strip()

            # Check for Roman numerals
            elif roman_pattern.match(line_stripped):
                is_section = True
                level = 1
                title = roman_pattern.match(line_stripped).group(1).strip()

            # Check for all caps (likely heading)
            elif line_stripped.isupper() and len(line_stripped) > 3 and len(line_stripped) < 100:
                is_section = True
                level = 1

            # Check for section keywords
            elif any(keyword in line_stripped.upper() for keyword in section_keywords):
                is_section = True
                level = 1

            if is_section:
                # Find section content (until next section or end)
                section_content = []
                section_start = char_offset
                section_end = char_offset + len(line)

                for j in range(i + 1, len(lines)):
                    next_line = lines[j].strip()
                    # Stop at next section
                    if (next_line.isupper() and len(next_line) > 3) or numbered_pattern.match(next_line):
                        break
                    section_content.append(lines[j])
                    section_end += len(lines[j]) + 1

                content = '\n'.join(section_content).strip()

                # Determine page number if page_map provided
                page_num = None
                if page_map:
                    for page, bounds in page_map.items():
                        if bounds['start'] <= section_start < bounds['end']:
                            page_num = page
                            break

                sections.append(Section(
                    level=level,
                    title=title,
                    start_char=section_start,
                    end_char=section_end,
                    content=content,
                    page=page_num
                ))

            char_offset += len(line) + 1

        return sections

    def extract_outline_from_html(self, soup: BeautifulSoup) -> List[Section]:
        """Extract outline from HTML headings"""
        sections = []
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

        for heading in headings:
            level = int(heading.name[1])  # h1 -> 1, h2 -> 2, etc.
            title = heading.get_text(strip=True)

            # Get content until next heading
            content_parts = []
            for sibling in heading.find_next_siblings():
                if sibling.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    break
                text = sibling.get_text(strip=True)
                if text:
                    content_parts.append(text)

            content = '\n'.join(content_parts)

            sections.append(Section(
                level=level,
                title=title,
                start_char=0,  # HTML doesn't have char positions
                end_char=len(content),
                content=content
            ))

        return sections

    def generate_snippets(self, text: str, num_snippets: int = 5, snippet_length: int = 200) -> List[str]:
        """
        Generate text snippets for preview

        Args:
            text: Full document text
            num_snippets: Number of snippets to generate
            snippet_length: Characters per snippet

        Returns:
            List of text snippets
        """
        snippets = []
        text_length = len(text)

        if text_length <= snippet_length:
            return [text]

        # Distribute snippets evenly through document
        step = text_length // num_snippets

        for i in range(num_snippets):
            start = i * step
            end = min(start + snippet_length, text_length)

            # Try to break at word boundary
            snippet = text[start:end]
            if end < text_length:
                last_space = snippet.rfind(' ')
                if last_space > snippet_length * 0.7:  # Only if reasonably close to end
                    snippet = snippet[:last_space] + '...'

            snippets.append(snippet.strip())

        return snippets

    def _calculate_confidence(self, text: str, warnings: List[str], pages: int = 1) -> float:
        """
        Calculate extraction confidence score

        Factors:
        - Text length (longer = more confident)
        - Number of warnings
        - Text quality (character distribution)
        """
        confidence = 1.0

        # Penalize for warnings
        confidence -= len(warnings) * 0.1

        # Penalize for very short text
        if len(text) < 100:
            confidence -= 0.3
        elif len(text) < 500:
            confidence -= 0.1

        # Penalize for low text/page ratio (likely scanned PDFs)
        if pages > 1:
            chars_per_page = len(text) / pages
            if chars_per_page < 100:
                confidence -= 0.4
            elif chars_per_page < 500:
                confidence -= 0.2

        # Check for high ratio of non-alphanumeric characters
        alphanum_ratio = sum(c.isalnum() or c.isspace() for c in text) / max(len(text), 1)
        if alphanum_ratio < 0.7:
            confidence -= 0.2

        return max(0.0, min(1.0, confidence))
